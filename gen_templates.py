from pathlib import Path
import csv
from openpyxl import Workbook
from docx import Document

out_dir = Path("frontend/public/templates")
out_dir.mkdir(parents=True, exist_ok=True)

csv_path = out_dir / "question-import-template.csv"
rows = list(csv.reader(open(csv_path, "r", encoding="utf-8")))
header = rows[0]
sample = rows[1:]

# XLSX template
wb = Workbook()
ws = wb.active
ws.title = "Questions"
ws.append(header)
for r in sample:
    ws.append(r)
wb.save(out_dir / "question-import-template.xlsx")

# DOCX template (simple structured text)
doc = Document()
doc.add_paragraph("Question import template")
doc.add_paragraph("")
letters = ["A", "B", "C", "D", "E", "F"]

for r in sample:
    data = {header[i]: (r[i] if i < len(r) else "") for i in range(len(header))}

    doc.add_paragraph(f"Question: {data.get('question', '').strip()}")
    doc.add_paragraph(f"Type: {data.get('type', '').strip()}")
    doc.add_paragraph(f"Difficulty: {(data.get('difficulty', '').strip() or 'medium')}")

    tags = (data.get('tags', '') or '').strip()
    doc.add_paragraph(f"Tags: {tags}")

    explanation = (data.get('explanation', '') or '').strip()
    doc.add_paragraph(f"Explanation: {explanation}")

    correct = (data.get('correct_options', '') or '').strip()
    doc.add_paragraph(f"Correct: {correct}")

    for idx, letter in enumerate(letters, start=1):
        col = f"option_{idx}"
        val = (data.get(col, '') or '').strip()
        if val:
            doc.add_paragraph(f"{letter}. {val}")

    doc.add_paragraph("")

doc.save(out_dir / "question-import-template.docx")
print('generated')
