from docx import Document
from pathlib import Path
import csv

out_dir = Path("frontend/public/templates")
out_docx = out_dir / "question-import-template.docx"

csv_path = out_dir / "question-import-template.csv"
rows = list(csv.reader(csv_path.open("r", encoding="utf-8")))
header = rows[0]
# Normalize header keys: remove UTF-8 BOM if present
header_norm = [h.replace("\ufeff", "").strip().lower() for h in header]

sample = rows[1:]

letters = ["A", "B", "C", "D", "E", "F"]

# Build a map from normalized header -> original header index
idx_by_key = {}
for i, k in enumerate(header_norm):
    idx_by_key[k] = i


doc = Document()
doc.add_paragraph("Question import template")
doc.add_paragraph("")

for r in sample:
    # Access helper
    def get_col(col_name: str, default: str = "") -> str:
        idx = idx_by_key.get(col_name)
        if idx is None:
            return default
        return r[idx] if idx < len(r) else default

    question = (get_col("question") or "").strip()
    qtype = (get_col("type") or "").strip()
    difficulty = (get_col("difficulty") or "").strip() or "medium"
    tags = (get_col("tags") or "").strip()
    explanation = (get_col("explanation") or "").strip()
    correct = (get_col("correct_options") or "").strip()

    doc.add_paragraph(f"Question: {question}")
    doc.add_paragraph(f"Type: {qtype}")
    doc.add_paragraph(f"Difficulty: {difficulty}")
    doc.add_paragraph(f"Tags: {tags}")
    doc.add_paragraph(f"Explanation: {explanation}")
    doc.add_paragraph(f"Correct: {correct}")

    for idx, letter in enumerate(letters, start=1):
        col = f"option_{idx}"
        val = (get_col(col) or "").strip()
        if val:
            doc.add_paragraph(f"{letter}. {val}")

    doc.add_paragraph("")

doc.save(out_docx)
print("regenerated", out_docx)
